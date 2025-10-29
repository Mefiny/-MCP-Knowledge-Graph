"""
Word Parser Module
解析Word文档（.docx），提取文本和元数据
"""
import os
from typing import Dict, List
from pathlib import Path
from docx import Document
from loguru import logger


class WordParser:
    """Word文档解析器"""

    def __init__(self):
        logger.info("WordParser initialized")

    def parse(self, file_path: str) -> Dict:
        """
        解析Word文档

        Args:
            file_path: Word文件路径

        Returns:
            {
                "text": "完整文本",
                "paragraphs": [{"index": 0, "text": "段落内容", "style": ""}],
                "tables": [{"index": 0, "data": [[]]  }],
                "metadata": {"title": "", "author": "", ...}
            }
        """
        try:
            # 打开Word文档
            doc = Document(file_path)

            # 提取元数据
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "file_name": Path(file_path).name,
                "paragraph_count": len(doc.paragraphs)
            }

            # 提取段落
            paragraphs = []
            full_text = []

            for idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:  # 只保留非空段落
                    paragraphs.append({
                        "index": idx,
                        "text": text,
                        "style": para.style.name if para.style else "Normal"
                    })
                    full_text.append(text)

            # 提取表格
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                tables.append({
                    "index": table_idx,
                    "data": table_data,
                    "rows": len(table.rows),
                    "cols": len(table.columns)
                })

            result = {
                "text": "\n\n".join(full_text),
                "paragraphs": paragraphs,
                "tables": tables,
                "metadata": metadata,
                "status": "success"
            }

            logger.info(f"Successfully parsed Word: {file_path}, {len(paragraphs)} paragraphs, {len(tables)} tables")
            return result

        except Exception as e:
            logger.error(f"Error parsing Word {file_path}: {str(e)}")
            return {
                "text": "",
                "paragraphs": [],
                "tables": [],
                "metadata": {},
                "status": "error",
                "error": str(e)
            }

    def extract_headings(self, file_path: str) -> List[Dict]:
        """提取所有标题（Heading样式）"""
        try:
            doc = Document(file_path)
            headings = []

            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    headings.append({
                        "level": para.style.name,
                        "text": para.text.strip()
                    })

            return headings
        except Exception as e:
            logger.error(f"Error extracting headings: {str(e)}")
            return []


# 测试代码
if __name__ == "__main__":
    parser = WordParser()

    # 测试文件
    test_file = "test.docx"
    if os.path.exists(test_file):
        result = parser.parse(test_file)
        print(f"Title: {result['metadata']['title']}")
        print(f"Author: {result['metadata']['author']}")
        print(f"Paragraphs: {result['metadata']['paragraph_count']}")
        print(f"Tables: {len(result['tables'])}")
        print(f"First 200 chars: {result['text'][:200]}...")
    else:
        print(f"Test file {test_file} not found")
