#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""读取上传文档的内容"""
import sys
import os
sys.path.insert(0, 'backend')

from docx import Document

# 读取第一个文档
doc_path = 'backend/uploads/5730718d-88de-4d46-94cd-5f032f21dfaf.docx'

try:
    doc = Document(doc_path)

    print("=" * 60)
    print("文档内容预览:")
    print("=" * 60)

    # 获取所有段落
    for i, para in enumerate(doc.paragraphs[:20]):
        text = para.text.strip()
        if text:
            # 安全输出，替换无法编码的字符
            safe_text = text.encode('utf-8', errors='replace').decode('utf-8')
            print(f"\n段落 {i+1}:")
            print(safe_text[:300])  # 只显示前300个字符

            # 检查是否包含"铠图"或"电子商务"
            if '铠图' in text or '电子商务' in text:
                print(f"\n>>> 发现关键词: 铠图/电子商务")

    # 统计
    total_text = '\n'.join([p.text for p in doc.paragraphs])
    print("\n" + "=" * 60)
    print("统计信息:")
    print("=" * 60)
    print(f"总段落数: {len(doc.paragraphs)}")
    print(f"总字符数: {len(total_text)}")
    print(f"包含'铠图': {'是' if '铠图' in total_text else '否'}")
    print(f"包含'电子商务': {'是' if '电子商务' in total_text else '否'}")

except Exception as e:
    print(f"错误: {e}")
    import traceback
    traceback.print_exc()
